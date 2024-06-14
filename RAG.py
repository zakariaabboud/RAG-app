import time


from langchain.text_splitter import RecursiveCharacterTextSplitter
import concurrent.futures
import docx
import openai
# import cohere
import numpy as np
import json
from decouple import config

instructions = """Tu es un expert des produits financiers de Cyrus.
Voici la liste produits Cyrus par catégorie : 
- SCPI : Pierval Santé, Transitions Europe, LOG IN.
- PER : PER Amytis Retraite, PER Generali Patrimoine 
- assurance vie : Spirica Initial, Spirica Perspective 8, Amytis Select 2, Multisupport excellence, General Espace Lux Vie
- Compte titres : Alpheys Invest
- Produit Structuré : Obligation Callable Taux Fixe Mai 2024
- Fond immobilier : Alcyon 2
- Club deal : Doc City, HLOG 2
- Non coté : Capza Rémploi, Ring Altitude Growth II, Eurazeo strategic opportunity 3, Mid Market Monde 2023
"""



#ranker = Ranker(model_name="ms-marco-MultiBERT-L-12")
Client = openai.OpenAI(api_key=config("openai_api"))
# co = cohere.Client(api_key="RsLI9vFKpeEUpZ787jK0h30jN8VlX32ABrqjB0NJ")
#Mistral_client = MistralClient(api_key=config("mistral_api"))


# def read_file(filenames):
#     """Read the pdf file and return the text."""
#     text = ""
#     for filename in filenames:
#         pdf = pypdfium.PdfDocument(filename)
#         for page in pdf:
#             text_page = page.get_textpage().get_text_range().replace('\r\n',' ')
#             text += text_page
#         pdf.close()
#     return text

def read_file(filenames):
    """Read the docx file and return the text."""
    text = ""
    for filename in filenames:
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text


def chunk(text):
    """Chunk the text into sentences and return a list of sentences."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=250,
        separators= ["\n\n","\n","."," ",""])
    return splitter.split_text(text)
        

def openai_embedding(text):
    """Get the embedding of the text using openai."""

    return Client.embeddings.create(model="text-embedding-3-small", input=text).data[0].embedding

def get_embeddings(model,sentences):
    """Get the embeddings of the sentences."""
    
    with concurrent.futures.ThreadPoolExecutor() as executor:
        embeddings = executor.map(model.encode, sentences)
    #embeddings = model.encode(sentences)
        
    return list(embeddings)

def get_embeddings_openai(sentences):
    """Get the embeddings of the sentences using openai."""
    
    with concurrent.futures.ThreadPoolExecutor() as executor:
        embeddings = executor.map(openai_embedding, sentences)
        
    return list(embeddings)

def prepare_chunks(filename):
    """Read the file and prepare the chunks."""

    file_dir = "chunks.json"
    text = read_file(filename)
   
    sentences = chunk(text)
    embeddings = get_embeddings_openai(sentences)

    
    data = {"embeddings":{},"sentences":[]}
    

    for i in range(len(sentences)):
        data["embeddings"][sentences[i]] = embeddings[i]
        data["sentences"].append(sentences[i])

    with open(file_dir, 'w') as outfile:
        json.dump(data, outfile)
    

def get_similarity(embeddings_1,embeddings_2):
    """Get the cosine similarity between two embeddings."""
    similarity = np.dot(embeddings_1,embeddings_2.T) / (np.linalg.norm(embeddings_1) * np.linalg.norm(embeddings_2,axis=1))
    return similarity
    

# def rerank(query, chunks, n = 5):
#     """Rerank the chunks."""
#     passages = [{"id": i, "text": chunks[i]} for i in range(len(chunks))]

#     request = RerankRequest(query=query, passages=passages)
#     response = ranker.rerank(request)
#     return response[:n]

# def cohere_rerank(query, chunks, n = 5):
#     """Rerank the chunks using cohere."""
#     results = []

#     response = co.rerank(query=query, documents=chunks, top_n = n, model="rerank-multilingual-v3.0")
#     for i in response.results:
#         results.append(i.index)

#     return results


def get_near_chunks(text,n = 5):
    """Get the chunks that are the most similar to the text."""
    
    file_dir = "chunks.json"
    with open(file_dir) as json_file:
        data = json.load(json_file)

    embeddings_1 = get_embeddings_openai([text])
    embeddings_2 = np.array(list(data["embeddings"].values()))

    similarity = get_similarity(embeddings_1,embeddings_2)

    # 15 nearest chunks
    index = np.argsort(similarity[0])[-1:-15-1:-1]
    
    # chunks_index = { data["sentences"][i]:i for i in index }
    

    # nearest 5 chunks with 2 chunks before and 2 chunks after
    
    near_chunks = [ " ".join(data["sentences"][ i - 1 : i +2]) for i in index[:n] ]

    # near_chunks = [ "".join(data["sentences"][ chunks_index[chunk] - 2 : chunks_index[chunk] + 3 ])
    #                  for chunk in near_chunks ]
    


    #near_chunks = [ data["sentences"][ chunks_index[i["text"]] ] for i in near_passages ]

    #near_chunks = [ data["sentences"][i] for i in index ]

    return near_chunks

def get_completion(messages, model="gpt-3.5-turbo-0125",temperature=0):
    
    response = Client.chat.completions.create(
        messages=messages,
        model=model,
        temperature=temperature,
    )

    return response.choices[0].message.content
# def get_completion_mistral(messages, model="mistral-small",temperature=0):
        
#         response = Mistral_client.chat(
#             messages=messages,
#             model=model,
#             temperature=temperature,
#         )
    
#         return response.choices[0].message.content

def decompose_question(question):
    """Decompose the question into sentences."""
    prompt = """Décompose la question en plusieurs requêtes de recherche plus facile et complémetaire. Les requêtes doivent être séparées par un retour à la ligne et ne doivent pas être numérotées.
    Évite de créer des requêtes similaires ou redondantes.

    SCPI : Pierval Santé, Transitions Europe, LOG IN.
    PER : PER Amytis Retraite, PER Generali Patrimoine.
    Assurance vie : Spirica Initial, Spirica Perspective 8, Amytis Select 2, Multisupport excellence, General Espace Lux Vie.
    Compte titres : Alpheys Invest.
    Produit Structuré : Obligation Callable Taux Fixe Mai 2024.
    Fond immobilier : Alcyon 2.
    Club deal : Doc City, HLOG 2.
    Non coté : Capza Rémploi, Ring Altitude Growth II, Eurazeo Strategic Opportunity 3, Mid Market Monde 2023.

    IMPORTANT : si la question concerne une catégorie de produits financiers de Cyrus, décompose-la en plusieurs requêtes de recherche pour chaque produit de la catégorie.
    Exemple : Question : Quels sont les avantages du PER ? Resultat : Quels sont les avantages du PER Amytis Retraite ? Quels sont les avantages du PER Generali Patrimoine ?
    """

    messages = [{"role": "system", "content" : prompt}]
    messages.append({"role": "user", "content" : question})
    #messages = [ChatMessage(role = "user", content = prompt)]

    gpt_respond = get_completion(messages)
    if gpt_respond[-1] == "\n":
        gpt_respond = gpt_respond[:-1]
    #gpt_respond = get_completion_mistral(messages)

    questions = gpt_respond.split("\n")

    print(questions)

    return questions

def get_answer(question,context,temperature=0.7,insutructions=""):
    """Get the answer to the question. The context is a list of sentences."""

    messages = []
    prompt = instructions
    prompt += "Répond à la question en utilisant le contexte suivant :  \n"
    for i in context:
        prompt += i + "\n"
    prompt += "La réponse doit être présice, concise et détaillé .\n"
    messages.append({"role": "system", "content" : prompt})
    messages.append({"role": "user", "content" : question})
    #messages.append(ChatMessage(role = "system", content = prompt))
    #messages.append(ChatMessage(role = "user", content = question))
    
    gpt_respond = get_completion(messages,temperature=temperature)
    return gpt_respond




def respond_1(question):
    """Get the answer to the question uitilizing the context of the first 4 chunks."""
    
    context = get_near_chunks(question)
    answer = get_answer(question,context)
    
    return answer

def respond_2(question):
    """Get the answer to the question with 6 chunks. 3 chunks at a time."""
    
    context = get_near_chunks(question)
    
    
    with concurrent.futures.ThreadPoolExecutor() as executor:
        query_1 = executor.submit(get_answer,question,context[:3])
        query_2 = executor.submit(get_answer,question,context[3:])

        answer_1 = query_1.result()
        answer_2 = query_2.result()
    
    
    context = [answer_1,answer_2]
    answer = get_answer(question,context)
    
    return answer



def respond_3(question):
    """Get the answer to the question by decomposing the question and using the context of the first 4 chunks for each question."""
    questions = decompose_question(question)
    
    with concurrent.futures.ThreadPoolExecutor() as executor:
        answers = executor.map(respond_1,questions)

    context = list(answers)
    answer = get_answer(question,context,temperature=1,insutructions=instructions)
    return answer


    






